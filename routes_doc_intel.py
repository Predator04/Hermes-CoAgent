"""Document intelligence routes.

Endpoints:
  POST /doc/extract - structured JSON extraction from PDFs, Office docs, and
                      scanned images. Returns per-page text, metadata, and
                      any embedded tables that can be recovered.
  POST /doc/table   - table-only extraction (rows/cells) for PDFs, XLSX
                      workbooks, and DOCX tables.

Input (both endpoints, JSON body):
  {"path": "C:/x.pdf"}                   local file path, or
  {"data": "<base64>", "filename": "x"}  inline bytes + filename hint
  optional: {"pages": "1-3,5"}           limit extraction to a page range
  optional: {"ocr": true|false|"auto"}   force / disable OCR fallback (default "auto")

Text extraction chain per type:
  .pdf                pdfplumber -> pypdf (text) ; pdfplumber (tables) ;
                     rasterize + reuse routes_ocr._windows_ocr when a page
                     has no extractable text ("auto" OCR)
  .docx / .doc        python-docx paragraphs + tables
  .xlsx / .xlsm       openpyxl sheets as tables
  .png / .jpg / .tiff routes_ocr._windows_ocr (Windows Media OCR)

All heavy deps (pypdf, pdfplumber, python-docx, openpyxl, PIL, routes_ocr)
are imported inside try/except so the Linux syntax-check CI stays green.
"""

import base64
import os
import re
import tempfile

from flask import jsonify

from shared import _json_body, _log, _missing_field


# ---------------------------------------------------------------------------
# Optional dependency probes. Everything below is guarded so a missing
# module never breaks import.
# ---------------------------------------------------------------------------
try:
    import pypdf  # type: ignore
    _HAS_PYPDF = True
except Exception:
    pypdf = None
    _HAS_PYPDF = False

try:
    import pdfplumber  # type: ignore
    _HAS_PDFPLUMBER = True
except Exception:
    pdfplumber = None
    _HAS_PDFPLUMBER = False

try:
    import docx as _docx  # python-docx
    _HAS_DOCX = True
except Exception:
    _docx = None
    _HAS_DOCX = False

try:
    from openpyxl import load_workbook as _load_workbook  # type: ignore
    _HAS_OPENPYXL = True
except Exception:
    _load_workbook = None
    _HAS_OPENPYXL = False

try:
    from PIL import Image as _PILImage  # type: ignore
    _HAS_PIL = True
except Exception:
    _PILImage = None
    _HAS_PIL = False


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
_MAX_INLINE_BYTES = 32 * 1024 * 1024  # 32 MB cap for base64 payloads
_MAX_PAGES = 500
_PAGE_RANGE_RE = re.compile(r"^\s*\d+\s*(-\s*\d+\s*)?$")

_PDF_SUFFIXES = (".pdf",)
_DOCX_SUFFIXES = (".docx", ".docm")
_DOC_LEGACY_SUFFIXES = (".doc",)
_XLSX_SUFFIXES = (".xlsx", ".xlsm")
_IMAGE_SUFFIXES = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp")


# ---------------------------------------------------------------------------
# Input resolution
# ---------------------------------------------------------------------------
def _resolve_source(body):
    """Return (path_or_None, tempfile_path_or_None, filename, error_response).

    On success error_response is None. Caller is responsible for unlinking
    the tempfile_path if it is not None.
    """
    path = body.get("path") or body.get("file")
    data_b64 = body.get("data")
    filename = (body.get("filename") or "").strip()

    if path:
        path = str(path)
        if not os.path.isfile(path):
            return None, None, None, (jsonify({"error": f"file not found: {path}"}), 404)
        return path, None, filename or os.path.basename(path), None

    if data_b64:
        if not filename:
            return None, None, None, (
                jsonify({"error": "filename required when using base64 data"}),
                400,
            )
        try:
            raw = base64.b64decode(str(data_b64), validate=False)
        except Exception as exc:
            return None, None, None, (
                jsonify({"error": f"invalid base64: {exc}"}),
                400,
            )
        if len(raw) > _MAX_INLINE_BYTES:
            return None, None, None, (
                jsonify({
                    "error": "payload too large",
                    "detail": f"{len(raw)} > {_MAX_INLINE_BYTES}",
                }),
                413,
            )
        suffix = os.path.splitext(filename)[1] or ".bin"
        fd, tmp_path = tempfile.mkstemp(suffix=suffix, prefix="doc_intel_")
        try:
            with os.fdopen(fd, "wb") as fh:
                fh.write(raw)
        except OSError as exc:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
            return None, None, None, (
                jsonify({"error": f"tempfile write failed: {exc}"}),
                500,
            )
        return tmp_path, tmp_path, filename, None

    return None, None, None, (
        jsonify({"error": "provide 'path', 'file', or 'data' (base64 with 'filename')"}),
        400,
    )


def _classify(filename):
    low = (filename or "").lower()
    if low.endswith(_PDF_SUFFIXES):
        return "pdf"
    if low.endswith(_DOCX_SUFFIXES):
        return "docx"
    if low.endswith(_DOC_LEGACY_SUFFIXES):
        return "doc"
    if low.endswith(_XLSX_SUFFIXES):
        return "xlsx"
    if low.endswith(_IMAGE_SUFFIXES):
        return "image"
    return "unknown"


def _parse_page_range(spec, total):
    """Parse a page spec like '1-3,5,7-9' into a sorted list of 1-based ints.

    Returns None on invalid input, or an empty list if nothing matches.
    total is used to clamp the upper bound.
    """
    if spec is None or not str(spec).strip():
        return list(range(1, min(total, _MAX_PAGES) + 1))
    if isinstance(spec, int):
        return [spec] if 1 <= spec <= min(total, _MAX_PAGES) else []
    if not isinstance(spec, str):
        return None
    pages = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if not _PAGE_RANGE_RE.match(part):
            return None
        if "-" in part:
            a, b = part.split("-", 1)
            try:
                lo = max(1, int(a.strip()))
                hi = min(total, int(b.strip()))
            except ValueError:
                return None
            if lo <= hi:
                pages.update(range(lo, hi + 1))
        else:
            try:
                n = int(part)
            except ValueError:
                return None
            if 1 <= n <= total:
                pages.add(n)
    return sorted(pages)[:_MAX_PAGES]


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------
def _pdf_page_count(path):
    if _HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(path) as pdf:
                return len(pdf.pages)
        except Exception:
            pass
    if _HAS_PYPDF:
        try:
            reader = pypdf.PdfReader(path)
            return len(reader.pages)
        except Exception:
            pass
    return 0


def _pdf_extract(path, page_spec, want_tables, ocr_mode):
    """Return (result_dict, error_or_none)."""
    if not (_HAS_PDFPLUMBER or _HAS_PYPDF):
        return None, "no PDF backend installed (pip install pdfplumber pypdf)"

    total = _pdf_page_count(path)
    if total == 0:
        return None, "unable to open PDF (0 pages)"
    pages_wanted = _parse_page_range(page_spec, total)
    if pages_wanted is None:
        return None, "invalid pages spec (use '1-3,5')"

    metadata = {}
    pages_out = []
    tables_out = []
    engines = set()

    # Metadata via pypdf when available (pdfplumber does not expose it uniformly).
    if _HAS_PYPDF:
        try:
            reader = pypdf.PdfReader(path)
            meta = reader.metadata or {}
            for k, v in (meta.items() if hasattr(meta, "items") else []):
                key = str(k).lstrip("/")
                try:
                    metadata[key] = str(v)
                except Exception:
                    continue
        except Exception as exc:
            _log(f"doc/extract pypdf metadata failed: {type(exc).__name__}: {exc}")

    # Text + tables via pdfplumber (preferred; better table support).
    text_by_page = {}
    if _HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(path) as pdf:
                engines.add("pdfplumber")
                for pno in pages_wanted:
                    page = pdf.pages[pno - 1]
                    try:
                        text_by_page[pno] = (page.extract_text() or "").strip()
                    except Exception as exc:
                        _log(f"doc/extract pdfplumber text p{pno} failed: {exc}")
                        text_by_page[pno] = ""
                    if want_tables:
                        try:
                            for tbl in page.extract_tables() or []:
                                if tbl:
                                    tables_out.append({"page": pno, "rows": tbl})
                        except Exception as exc:
                            _log(f"doc/extract pdfplumber tables p{pno} failed: {exc}")
        except Exception as exc:
            _log(f"doc/extract pdfplumber open failed: {type(exc).__name__}: {exc}")

    # Fill missing pages with pypdf text.
    if _HAS_PYPDF and any(pno not in text_by_page or not text_by_page[pno] for pno in pages_wanted):
        try:
            reader = pypdf.PdfReader(path)
            engines.add("pypdf")
            for pno in pages_wanted:
                if text_by_page.get(pno):
                    continue
                try:
                    text_by_page[pno] = (reader.pages[pno - 1].extract_text() or "").strip()
                except Exception as exc:
                    _log(f"doc/extract pypdf text p{pno} failed: {exc}")
                    text_by_page.setdefault(pno, "")
        except Exception as exc:
            _log(f"doc/extract pypdf open failed: {type(exc).__name__}: {exc}")

    # OCR fallback for scan-only pages.
    ocr_pages = []
    if ocr_mode in ("auto", True, "true", "1"):
        empty_pages = [p for p in pages_wanted if not text_by_page.get(p)]
        if empty_pages:
            ocr_text_by_page = _pdf_ocr_pages(path, empty_pages)
            for pno, txt in ocr_text_by_page.items():
                if txt:
                    text_by_page[pno] = txt
                    ocr_pages.append(pno)
            if ocr_text_by_page:
                engines.add("windows_ocr")

    all_text_parts = []
    for pno in pages_wanted:
        txt = text_by_page.get(pno, "")
        pages_out.append({"page": pno, "text": txt, "chars": len(txt)})
        if txt:
            all_text_parts.append(txt)

    result = {
        "type": "pdf",
        "engine": sorted(engines) or ["none"],
        "page_count": total,
        "pages": pages_out,
        "text": "\n\n".join(all_text_parts),
        "metadata": metadata,
    }
    if ocr_pages:
        result["ocr_pages"] = ocr_pages
    if want_tables:
        result["tables"] = tables_out
        result["table_count"] = len(tables_out)
    return result, None


def _pdf_ocr_pages(path, page_numbers):
    """Rasterize the given 1-based pages via pdfplumber and OCR each one.

    Returns a dict {page_number: text}. Silently skips pages that cannot be
    rendered or OCRed - this is a best-effort fallback.
    """
    if not (_HAS_PDFPLUMBER and _HAS_PIL):
        return {}
    try:
        from routes_ocr import _windows_ocr  # type: ignore
    except Exception as exc:
        _log(f"doc/extract OCR fallback unavailable: {type(exc).__name__}: {exc}")
        return {}

    out = {}
    try:
        with pdfplumber.open(path) as pdf:
            for pno in page_numbers:
                try:
                    page = pdf.pages[pno - 1]
                    page_image = page.to_image(resolution=200)
                    pil = page_image.original
                    try:
                        res = _windows_ocr(pil)
                        if isinstance(res, dict) and res.get("success"):
                            out[pno] = (res.get("text") or "").strip()
                    finally:
                        try:
                            pil.close()
                        except Exception:
                            pass
                except Exception as exc:
                    _log(f"doc/extract OCR page {pno} failed: {type(exc).__name__}: {exc}")
    except Exception as exc:
        _log(f"doc/extract OCR pdfplumber open failed: {type(exc).__name__}: {exc}")
    return out


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------
def _docx_extract(path, want_tables):
    if not _HAS_DOCX:
        return None, "python-docx not installed (pip install python-docx)"
    try:
        doc = _docx.Document(path)
    except Exception as exc:
        return None, f"open failed: {type(exc).__name__}: {exc}"

    paragraphs = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            paragraphs.append(text)

    tables_out = []
    for idx, tbl in enumerate(doc.tables):
        rows = []
        for row in tbl.rows:
            rows.append([(cell.text or "").strip() for cell in row.cells])
        if rows:
            tables_out.append({"index": idx, "rows": rows})

    core = getattr(doc, "core_properties", None)
    metadata = {}
    if core is not None:
        for attr in ("title", "author", "subject", "keywords", "last_modified_by"):
            val = getattr(core, attr, None)
            if val:
                metadata[attr] = str(val)
        for attr in ("created", "modified"):
            val = getattr(core, attr, None)
            if val:
                try:
                    metadata[attr] = val.isoformat()
                except Exception:
                    metadata[attr] = str(val)

    result = {
        "type": "docx",
        "engine": ["python-docx"],
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "text": "\n".join(paragraphs),
        "metadata": metadata,
    }
    if want_tables:
        result["tables"] = tables_out
        result["table_count"] = len(tables_out)
    return result, None


# ---------------------------------------------------------------------------
# XLSX extraction
# ---------------------------------------------------------------------------
def _xlsx_extract(path, want_tables):
    if not _HAS_OPENPYXL:
        return None, "openpyxl not installed (pip install openpyxl)"
    try:
        wb = _load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:
        return None, f"open failed: {type(exc).__name__}: {exc}"

    sheets_out = []
    tables_out = []
    total_cells = 0
    try:
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = []
                for v in row:
                    if v is None:
                        cells.append("")
                    else:
                        cells.append(v if isinstance(v, (int, float, bool)) else str(v))
                # Drop trailing all-empty rows to keep payload small.
                if any(c not in ("", None) for c in cells):
                    rows.append(cells)
                    total_cells += len(cells)
            sheets_out.append({
                "name": sheet.title,
                "rows": rows,
                "row_count": len(rows),
            })
            if want_tables and rows:
                tables_out.append({"sheet": sheet.title, "rows": rows})
    finally:
        try:
            wb.close()
        except Exception:
            pass

    text_parts = []
    for s in sheets_out:
        text_parts.append(f"# {s['name']}")
        for row in s["rows"]:
            text_parts.append("\t".join(str(c) for c in row))
    result = {
        "type": "xlsx",
        "engine": ["openpyxl"],
        "sheets": sheets_out,
        "sheet_count": len(sheets_out),
        "cell_count": total_cells,
        "text": "\n".join(text_parts),
        "metadata": {},
    }
    if want_tables:
        result["tables"] = tables_out
        result["table_count"] = len(tables_out)
    return result, None


# ---------------------------------------------------------------------------
# Image (scan) OCR
# ---------------------------------------------------------------------------
def _image_extract(path):
    if not _HAS_PIL:
        return None, "Pillow not installed (pip install pillow)"
    try:
        from routes_ocr import _windows_ocr  # type: ignore
    except Exception as exc:
        return None, f"routes_ocr unavailable: {type(exc).__name__}: {exc}"
    try:
        pil = _PILImage.open(path)
    except Exception as exc:
        return None, f"image open failed: {type(exc).__name__}: {exc}"

    try:
        res = _windows_ocr(pil)
    finally:
        try:
            pil.close()
        except Exception:
            pass
    if not isinstance(res, dict) or not res.get("success"):
        detail = (res or {}).get("error") if isinstance(res, dict) else "unknown"
        return None, f"OCR failed: {detail}"
    text = (res.get("text") or "").strip()
    return {
        "type": "image",
        "engine": ["windows_ocr"],
        "text": text,
        "word_count": res.get("word_count", 0),
        "line_count": res.get("line_count", 0),
        "words": res.get("words", []),
        "metadata": {},
    }, None


# ---------------------------------------------------------------------------
# Flask registration
# ---------------------------------------------------------------------------
def register_routes(app, state, require_auth):

    @app.route("/doc/extract", methods=["POST"])
    @require_auth
    def route_doc_extract():
        body = _json_body() or {}
        source_path, tmp_path, filename, err = _resolve_source(body)
        if err:
            return err
        page_spec = body.get("pages")
        ocr_raw = body.get("ocr", "auto")
        if isinstance(ocr_raw, str):
            ocr_mode = ocr_raw.strip().lower()
        elif isinstance(ocr_raw, bool):
            ocr_mode = "true" if ocr_raw else "false"
        else:
            ocr_mode = "auto"

        try:
            try:
                kind = _classify(filename)
                if kind == "pdf":
                    result, err = _pdf_extract(source_path, page_spec, want_tables=True, ocr_mode=ocr_mode)
                elif kind == "docx":
                    result, err = _docx_extract(source_path, want_tables=True)
                elif kind == "doc":
                    result, err = None, "legacy .doc not supported (convert to .docx)"
                elif kind == "xlsx":
                    result, err = _xlsx_extract(source_path, want_tables=True)
                elif kind == "image":
                    result, err = _image_extract(source_path)
                else:
                    result, err = None, f"unsupported file type: {filename}"
            except Exception as exc:
                _log(f"doc/extract exception for {filename}: {type(exc).__name__}: {exc}")
                return jsonify({
                    "ok": False,
                    "error": f"extraction failed: {type(exc).__name__}: {exc}",
                    "filename": filename,
                }), 500

            if err:
                _log(f"doc/extract failed for {filename}: {err}")
                return jsonify({
                    "ok": False,
                    "error": err,
                    "filename": filename,
                    "type": kind,
                }), 400 if kind == "unknown" else 500

            result["ok"] = True
            result["filename"] = filename
            _log(f"doc/extract ok filename={filename!r} type={result.get('type')}")
            return jsonify(result)
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

    @app.route("/doc/table", methods=["POST"])
    @require_auth
    def route_doc_table():
        body = _json_body() or {}
        source_path, tmp_path, filename, err = _resolve_source(body)
        if err:
            return err
        page_spec = body.get("pages")

        try:
            kind = _classify(filename)
            tables = []
            engine = "none"
            if kind == "pdf":
                if not _HAS_PDFPLUMBER:
                    return jsonify({
                        "ok": False,
                        "error": "pdfplumber not installed (pip install pdfplumber)",
                    }), 500
                total = _pdf_page_count(source_path)
                if total == 0:
                    return jsonify({"ok": False, "error": "unable to open PDF"}), 500
                pages_wanted = _parse_page_range(page_spec, total)
                if pages_wanted is None:
                    return jsonify({"ok": False, "error": "invalid pages spec"}), 400
                engine = "pdfplumber"
                try:
                    with pdfplumber.open(source_path) as pdf:
                        for pno in pages_wanted:
                            try:
                                for tbl in pdf.pages[pno - 1].extract_tables() or []:
                                    if tbl:
                                        tables.append({"page": pno, "rows": tbl})
                            except Exception as exc:
                                _log(f"doc/table pdfplumber p{pno} failed: {exc}")
                except Exception as exc:
                    return jsonify({
                        "ok": False,
                        "error": f"pdfplumber failed: {type(exc).__name__}: {exc}",
                    }), 500
            elif kind == "docx":
                result, err2 = _docx_extract(source_path, want_tables=True)
                if err2:
                    return jsonify({"ok": False, "error": err2}), 500
                tables = result.get("tables", [])
                engine = "python-docx"
            elif kind == "xlsx":
                result, err2 = _xlsx_extract(source_path, want_tables=True)
                if err2:
                    return jsonify({"ok": False, "error": err2}), 500
                tables = result.get("tables", [])
                engine = "openpyxl"
            else:
                return jsonify({
                    "ok": False,
                    "error": f"tables not supported for type: {kind}",
                    "filename": filename,
                }), 400

            _log(f"doc/table ok filename={filename!r} type={kind} tables={len(tables)}")
            return jsonify({
                "ok": True,
                "filename": filename,
                "type": kind,
                "engine": engine,
                "count": len(tables),
                "tables": tables,
            })
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
